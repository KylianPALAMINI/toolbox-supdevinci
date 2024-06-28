import os
import subprocess
import threading
import tkinter as tk
from tkinter import ttk, messagebox
from PIL import Image, ImageTk
import webbrowser
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

# suivre le mode sombre
dark_mode = False

# savoir si un programme est en cours
program_running = False

# définir les chemins pour les programmes, icônes et logos
base_dir = os.path.dirname(__file__)
programs_dir = os.path.join(base_dir, "programmes")
icons_dir = os.path.join(programs_dir, "icones")
logos_dir = os.path.join(base_dir, "programmes/logos")
temp_dir_root = os.path.join(base_dir, "temp")
temp_dir_programs = os.path.join(programs_dir, "temp")
general_report_dir = os.path.join(base_dir, "rapports/general")

def open_program(program_name):
    global program_running
    if program_running:
        messagebox.showwarning("Attention", "un programme est déjà en cours d'exécution.")
        return

    def run_program():
        global program_running
        program_running = True
        disable_buttons()
        program_path = os.path.join(programs_dir, program_name)
        print(f"lancement du programme : {program_path}")  # debug print
        subprocess.run(['python', program_path], creationflags=subprocess.CREATE_NO_WINDOW)
        program_running = False
        enable_buttons()

    # lancer un thread pour pas bloquer l'interface
    program_thread = threading.Thread(target=run_program)
    program_thread.start()

def disable_buttons():
    for button_frame in button_frames:
        for widget in button_frame.winfo_children():
            widget.config(state=tk.DISABLED)

def enable_buttons():
    for button_frame in button_frames:
        for widget in button_frame.winfo_children():
            widget.config(state=tk.NORMAL)

def toggle_mode():
    global dark_mode
    dark_mode = not dark_mode
    apply_theme()

def apply_theme():
    if dark_mode:
        root.config(bg="#2e2e2e")
        title_label.config(bg="#2e2e2e", fg="white")
        created_by_label.config(bg="#2e2e2e", fg="white")
        mode_button.config(bg="#444444", fg="white")
        for button_frame in button_frames:
            button_frame.config(bg="#444444")
            button_frame.pack_configure(pady=5)
            for widget in button_frame.winfo_children():
                widget.config(bg="#444444", fg="white", activebackground="#555555", activeforeground="white")
    else:
        root.config(bg="white")
        title_label.config(bg="white", fg="black")
        created_by_label.config(bg="white", fg="black")
        mode_button.config(bg="lightgray", fg="black")
        for button_frame in button_frames:
            button_frame.config(bg="lightgray")
            button_frame.pack_configure(pady=5)
            for widget in button_frame.winfo_children():
                widget.config(bg="lightgray", fg="black", activebackground="#dddddd", activeforeground="black")

def create_button(frame, text, icon_name, program_name):
    icon_path = os.path.join(logos_dir, icon_name)
    icon = Image.open(icon_path)
    icon = icon.resize((50, 50), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon)

    button_frame = tk.Frame(frame, bg="lightgray")
    button_frame.pack(fill=tk.X, pady=5)

    icon_label = tk.Label(button_frame, image=icon_photo, bg="lightgray")
    icon_label.image = icon_photo
    icon_label.pack(side=tk.LEFT, padx=10)

    button = tk.Button(button_frame, text=text, compound=tk.LEFT, command=lambda p=program_name: open_program(p), bg="lightgray")
    button.pack(side=tk.LEFT, fill=tk.X, expand=True)

    button_frames.append(button_frame)

def open_github():
    webbrowser.open("https://github.com/KylianPALAMINI/toolbox-supdevinci.git")

def show_about():
    about_text = (
        "Présentation Globale du Programme:\n\n"
        "Le programme 'Cybersécurité Toolbox' est une suite d'outils destinée aux professionnels et aux étudiants "
        "en cybersécurité. Ce programme offre une gamme d'utilitaires pour effectuer divers types de tests et d'analyses "
        "sur des réseaux et des systèmes informatiques, afin d'identifier des vulnérabilités, évaluer la sécurité et "
        "renforcer les défenses contre les cyberattaques.\n\n"
        "Situations d'Utilisation:\n"
        "Ce programme peut être utilisé dans de nombreuses situations, notamment:\n"
        "- Évaluation de la sécurité d'un réseau interne ou d'un réseau d'entreprise.\n"
        "- Tests de pénétration pour identifier les faiblesses potentielles avant qu'un attaquant ne les exploite.\n"
        "- Formation et éducation en cybersécurité pour les étudiants et les professionnels.\n"
        "- Validation de la configuration de la sécurité après des modifications ou des mises à jour.\n"
        "- Surveillance continue de la sécurité pour détecter les nouvelles vulnérabilités.\n\n"
        "Fonctionnalités Proposées:\n"
        "1. Scan de Ports:\n"
        "   - Effectue un scan des ports pour identifier les ports ouverts sur une cible spécifique.\n"
        "   - Utilise Nmap pour fournir des détails sur les services en cours d'exécution et leur état.\n"
        "   - Utile pour découvrir les points d'entrée potentiels pour les attaquants.\n\n"
        "2. Scan de Vulnérabilités:\n"
        "   - Utilise l'API Shodan pour rechercher les vulnérabilités connues sur une cible.\n"
        "   - Effectue également des scans de ports avec Scapy pour une analyse plus détaillée.\n"
        "   - Génère un rapport complet des vulnérabilités trouvées, y compris les descriptions et les scores CVSS.\n\n"
        "3. Bruteforce FTP:\n"
        "   - Tente de trouver le mot de passe FTP d'une cible en utilisant une liste de mots de passe.\n"
        "   - Utilise une approche multithread pour accélérer le processus de bruteforce.\n"
        "   - Idéal pour tester la robustesse des mots de passe FTP.\n\n"
        "4. Bruteforce SSH:\n"
        "   - Tente de trouver le mot de passe SSH d'une cible en utilisant une liste de mots de passe.\n"
        "   - Utilise une approche multithread pour accélérer le processus de bruteforce.\n"
        "   - Utile pour évaluer la sécurité des accès SSH.\n\n"
        "5. Extracteur de Données:\n"
        "   - Permet de télécharger des fichiers depuis une machine distante via SSH.\n"
        "   - Offre des fonctionnalités de scan de réseau pour identifier les hôtes actifs.\n"
        "   - Pratique pour l'analyse de fichiers sur des systèmes distants.\n\n"
        "6. Scan CVE:\n"
        "   - Recherche les vulnérabilités CVE associées aux services en cours d'exécution sur une cible.\n"
        "   - Utilise l'API Vulners pour obtenir des informations détaillées sur les CVE.\n"
        "   - Génère des rapports incluant des descriptions de vulnérabilités et des recommandations de mitigation.\n\n"
        "7. Test de Débit/Connexion:\n"
        "   - Effectue des tests de vitesse de connexion pour évaluer les débits de téléchargement et d'upload.\n"
        "   - Utilise des pings et des tests de vitesse pour fournir une évaluation complète de la qualité de la connexion.\n"
        "   - Idéal pour diagnostiquer les problèmes de réseau et de performance.\n\n"
        "Rappel Important:\n"
        "Ce programme est conforme aux recommandations de l'ANSSI (Agence Nationale de la Sécurité des Systèmes d'Information). "
        "Il est crucial de rappeler que l'utilisation de ce logiciel doit se faire dans un cadre légal et éthique. "
        "L'utilisation de ce programme à des fins malveillantes, telles que l'intrusion non autorisée ou la compromission de systèmes, "
        "est strictement interdite et passible de sanctions pénales conformément à l'article 323-3 du Code pénal français. "
        "Cet article stipule que le fait d'accéder frauduleusement ou de se maintenir dans tout ou partie d'un système de traitement automatisé "
        "de données est puni de deux ans d'emprisonnement et de 30 000 euros d'amende.\n\n"
        "Conclusion:\n"
        "La 'Cybersécurité Toolbox' est un ensemble d'outils puissants et polyvalents pour quiconque souhaite renforcer la sécurité de ses systèmes "
        "informatiques. En suivant les bonnes pratiques et en respectant les lois en vigueur, vous pouvez utiliser ces outils pour identifier "
        "et corriger les failles de sécurité, contribuant ainsi à un environnement numérique plus sûr."
    )
    messagebox.showinfo("À propos", about_text)

def clear_temp_folder():
    if os.path.exists(temp_dir_root):
        shutil.rmtree(temp_dir_root)
        os.makedirs(temp_dir_root)
    if os.path.exists(temp_dir_programs):
        shutil.rmtree(temp_dir_programs)
        os.makedirs(temp_dir_programs)
    messagebox.showinfo("Information", "Les dossiers temporaires ont été vidés.")

def generate_global_report():
    # Copier les fichiers de /programmes/temp vers le dossier racine /temp
    if os.path.exists(temp_dir_programs):
        for item in os.listdir(temp_dir_programs):
            s = os.path.join(temp_dir_programs, item)
            d = os.path.join(temp_dir_root, item)
            if os.path.isdir(s):
                shutil.copytree(s, d, dirs_exist_ok=True)
            else:
                shutil.copy2(s, d)

    if not os.path.exists(temp_dir_root) or not os.listdir(temp_dir_root):
        messagebox.showwarning("Attention", "Le dossier temporaire est vide.")
        return

    if not os.path.exists(general_report_dir):
        os.makedirs(general_report_dir)

    today = datetime.now().strftime("%Y%m%d")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_name = f"rapport_global_{timestamp}.docx"
    report_path = os.path.join(general_report_dir, report_name)
    pdf_report_path = os.path.join(general_report_dir, f"rapport_global_{timestamp}.pdf")

    doc = Document()
    doc.add_heading('Rapport Global', 0)

    for file_name in os.listdir(temp_dir_root):
        if file_name.endswith(".docx"):
            file_path = os.path.join(temp_dir_root, file_name)
            sub_doc = Document(file_path)
            doc.add_heading(file_name, level=1)
            # Inclure l'adresse cible dans le rapport global
            target_address = extract_target_address(sub_doc)
            if target_address:
                doc.add_paragraph(f"Adresse cible: {target_address}", style='Intense Quote')
            for element in sub_doc.element.body:
                doc.element.body.append(element)

    doc.save(report_path)
    
    # Convertir le rapport DOCX en PDF
    try:
        convert(report_path, pdf_report_path)
        messagebox.showinfo("Information", f"Rapport global généré : {report_path}\nRapport PDF généré : {pdf_report_path}")
    except Exception as e:
        messagebox.showerror("Erreur", f"Une erreur s'est produite lors de la conversion en PDF : {e}")

def extract_target_address(doc):
    """Extrait l'adresse cible du rapport individuel."""
    for para in doc.paragraphs:
        if para.text.startswith('Rapport de scan pour'):
            # Supposons que l'adresse soit la dernière partie du titre
            return para.text.split()[-1]
    return None

# configurer la fenetre principale
root = tk.Tk()
root.title("menu principal")
root.geometry("600x720")
root.iconbitmap(os.path.join(logos_dir, "menu_principal.ico"))

title_label = tk.Label(root, text="Cybersécurité Toolbox", font=("Helvetica", 24, "bold"))
title_label.pack(pady=20)

nav_frame = tk.Frame(root)
nav_frame.pack(pady=5)

# ajouter les boutons github et à propos
github_button = tk.Button(nav_frame, text="Github", command=open_github)
github_button.pack(side=tk.LEFT, padx=10)

about_button = tk.Button(nav_frame, text="À propos", command=show_about)
about_button.pack(side=tk.LEFT, padx=10)

clear_temp_button = tk.Button(nav_frame, text="Vider le dossier Temp", command=clear_temp_folder)
clear_temp_button.pack(side=tk.LEFT, padx=10)

generate_report_button = tk.Button(nav_frame, text="Générer Rapport Global", command=generate_global_report)
generate_report_button.pack(side=tk.LEFT, padx=10)

mode_button = tk.Button(root, text="Mode sombre", command=toggle_mode)
mode_button.pack(pady=10)

frame = tk.Frame(root)
frame.pack(pady=20)

button_frames = []

# ajouter les boutons pour chaque programme
programs = [
    ("Scan de ports", "scan_port.ico", "scan_port_v3_ui.py"),
    ("Scan de vulnérabilités", "scan_vuln.ico", "vulnera_scan_v1_ui.py"),
    ("Bruteforce ftp", "bruteforce_ftp.ico", "bruteforce_ftp_v1_ui.py"),
    ("Bruteforce ssh", "bruteforce_ssh.ico", "bruteforce_ssh_v1_ui.py"),
    ("Extracteur de données", "data_extractor.ico", "extracteur_data_v4_ui.py"),
    ("Scan cve", "scan_cve.ico", "scan_cve_01_ui.py"),
    ("Test de débit/connexion", "test_connection.ico", "test_debit_connexion_ui.py"),
]

for text, icon_name, program_name in programs:
    create_button(frame, text, icon_name, program_name)

created_by_label = tk.Label(root, text="Créé par kylian palamini, étudiant master 1 sup de vinci", font=("Helvetica", 10))
created_by_label.pack(side=tk.BOTTOM, pady=20)

apply_theme()
root.mainloop()
