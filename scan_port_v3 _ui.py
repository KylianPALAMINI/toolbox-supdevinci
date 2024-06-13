import nmap
import subprocess
from pythonping import ping
from docx import Document
from docx.shared import RGBColor
from datetime import datetime
from colorama import init, Fore, Style
import os as os_module
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext

init(autoreset=True)  # Initialise colorama pour réinitialiser automatiquement les styles après chaque print

def add_colored_text(paragraph, text, color):
    run = paragraph.add_run(text)
    run.font.color.rgb = color

# Dictionnaire des ports communs avec leurs descriptions
common_ports = {
    20: "FTP - File Transfer Protocol (Transfert de données)",
    21: "FTP - File Transfer Protocol (Contrôle)",
    22: "SSH - Secure Shell",
    23: "Telnet - Communications texte non sécurisées",
    25: "SMTP - Simple Mail Transfer Protocol",
    53: "DNS - Domain Name System",
    67: "DHCP - Dynamic Host Configuration Protocol (Serveur)",
    68: "DHCP - Dynamic Host Configuration Protocol (Client)",
    69: "TFTP - Trivial File Transfer Protocol",
    80: "HTTP - Hypertext Transfer Protocol",
    110: "POP3 - Post Office Protocol",
    123: "NTP - Network Time Protocol",
    137: "NetBIOS Name Service",
    138: "NetBIOS Datagram Service",
    139: "NetBIOS Session Service",
    143: "IMAP - Internet Message Access Protocol",
    161: "SNMP - Simple Network Management Protocol",
    162: "SNMP Trap",
    179: "BGP - Border Gateway Protocol",
    389: "LDAP - Lightweight Directory Access Protocol",
    443: "HTTPS - HTTP Secure",
    445: "Microsoft-DS Active Directory, partages Windows",
    465: "SMTPS - SMTP sécurisé",
    500: "IKE - Internet Key Exchange pour VPN",
    587: "SMTP sur TLS/SSL",
    636: "LDAPS - LDAP sur SSL",
    993: "IMAPS - IMAP sur SSL",
    995: "POP3S - POP3 sur SSL",
    1433: "Microsoft SQL Server management de base de données",
    1521: "Oracle base de données écouteur par défaut",
    1723: "PPTP - Point to Point Tunneling Protocol",
    2049: "NFS - Network File System",
    2082: "cPanel par défaut",
    2083: "cPanel sur SSL",
    2181: "Apache ZooKeeper port par défaut",
    2375: "Docker REST API (non chiffré)",
    2376: "Docker REST API (chiffré)",
    2483: "Oracle base de données écoute les connexions clients non sécurisées sur l'écouteur",
    2484: "Oracle base de données écoute les connexions clients SSL sur l'écouteur",
    3306: "MySQL - Service de base de données",
    3389: "RDP - Remote Desktop Protocol",
    5432: "PostgreSQL service de base de données",
    5672: "AMQP - Advanced Message Queuing Protocol",
    5900: "VNC - Virtual Network Computing",
    5984: "CouchDB service de base de données",
    6379: "Redis magasin de valeur clé",
    8080: "HTTP alternatif (souvent utilisé pour le proxy et la mise en cache)",
    9200: "Elasticsearch—port par défaut d'Elasticsearch",
    9300: "Elasticsearch—port de transport par défaut d'Elasticsearch",
    11211: "Service de mise en cache Memcached",
    27017: "MongoDB port par défaut"
}

def check_host_alive(host):
    """Vérifie si l'hôte est en vie en utilisant le ping."""
    try:
        response = ping(host, count=4, verbose=True)
        if response.success():
            return True
        else:
            return False
    except Exception as e:
        return False

def scan_ports(scanner, host, text_widget):
    """Effectue un scan des ports de l'hôte spécifié."""
    text_widget.insert(tk.END, f"Début du scan des ports pour {host}...\n", "info")
    text_widget.update()
    scanner.scan(host, '1-1024', arguments='-Pn')  # Scan des ports de 1 à 1024
    ports_info = []
    if host not in scanner.all_hosts():
        text_widget.insert(tk.END, f"Aucune information disponible pour l'hôte {host}\n", "warning")
        return []
    for port in scanner[host].all_tcp():
        state = scanner[host]['tcp'][port]['state']
        ports_info.append((port, state))
        text_widget.insert(tk.END, f"Le port {port} est {state}\n", "info")
    return ports_info

def scan_services_and_os(scanner, host, text_widget):
    """Effectue un scan des services et du système d'exploitation de l'hôte spécifié."""
    text_widget.insert(tk.END, f"Début du scan des services et OS pour {host}...\n", "info")
    text_widget.update()
    scanner.scan(host, arguments='-sV -O -Pn --osscan-guess --version-all --script banner')
    os_info = []
    services_info = []
    banners = []
    if host not in scanner.all_hosts():
        text_widget.insert(tk.END, f"Aucune information disponible pour l'hôte {host}\n", "warning")
        return [], [], []
    if 'osmatch' in scanner[host]:
        for osmatch in scanner[host]['osmatch']:
            if int(osmatch['accuracy']) >= 95:  # Filtre de précision de 95%
                os_info.append((osmatch['name'], osmatch['accuracy']))
                text_widget.insert(tk.END, f"Système d'exploitation possible : {osmatch['name']} (Précision : {osmatch['accuracy']}%)\n", "info")
    if 'tcp' in scanner[host]:
        for port in scanner[host]['tcp']:
            service = scanner[host]['tcp'][port]['name']
            services_info.append((port, service))
            text_widget.insert(tk.END, f"Service sur le port {port} : {service}\n", "info")
            if 'product' in scanner[host]['tcp'][port]:
                banners.append((port, scanner[host]['tcp'][port]['product']))
    return os_info[:3], services_info, banners  # Limiter à 3 résultats les plus précis

def generate_report_docx(host, ports_info, os_info, services_info, banners, folder):
    doc = Document()
    current_time = datetime.now().strftime("%d-%m-%Y %H:%M:%S")  # Format de date et heure

    title = doc.add_heading(level=0)
    add_colored_text(title, f'Rapport de scan pour {host}', RGBColor(0x42, 0x24, 0xE9))  # Bleu foncé
    # Ajout de la date et de l'heure
    header = doc.add_heading(level=2)
    add_colored_text(header, f'Généré le : {current_time}', RGBColor(0x77, 0x88, 0x99))  # Gris

    doc.add_heading('Scan des ports:', level=1)
    for port, state in ports_info:
        description = common_ports.get(port, "Service inconnu")
        p = doc.add_paragraph()
        add_colored_text(p, f'Port {port} ({description}): {state}', RGBColor(0x00, 0x8B, 0x8B))  # Vert foncé

    doc.add_heading("Supposition du système d'exploitation:", level=1)
    if not os_info:
        p = doc.add_paragraph()
        add_colored_text(p, 'Aucune information fiable sur le système d\'exploitation.', RGBColor(0xFF, 0x45, 0x00))  # Orange
    for os, accuracy in os_info:
        p = doc.add_paragraph()
        add_colored_text(p, f'{os} (Précision: {accuracy}%)', RGBColor(0xFF, 0x45, 0x00))  # Orange

    doc.add_heading('Services détectés:', level=1)
    for port, service in services_info:
        p = doc.add_paragraph()
        add_colored_text(p, f'Port {port}: {service}', RGBColor(0x6A, 0x5A, 0xCD))  # Slate blue

    doc.add_heading('Bannières collectées:', level=1)
    for port, banner in banners:
        p = doc.add_paragraph()
        add_colored_text(p, f'Port {port}: {banner}', RGBColor(0xDC, 0x14, 0x3C))  # Crimson

    os_module.makedirs(folder, exist_ok=True)  # Utilisation de l'alias pour créer le dossier
    filename = os_module.path.join(folder, f"rapport_scan_{host}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(filename)
    return filename

def start_scan():
    target_hosts = host_entry.get()
    if not target_hosts:
        messagebox.showwarning("Avertissement", "Veuillez entrer au moins un hôte.")
        return
    
    hosts = [host.strip() for host in target_hosts.split(',')]
    scanner = nmap.PortScanner()
    current_date = datetime.now().strftime('%d-%m-%Y')
    folder = f"rapport_programme_scan_{current_date}"
    results_text.delete(1.0, tk.END)
    results_text.insert(tk.END, "Début des scans...\n", "info")

    for host in hosts:
        results_text.insert(tk.END, f"Vérification si {host} est en ligne...\n", "info")
        results_text.update()
        if check_host_alive(host):
            results_text.insert(tk.END, f"{host} est en ligne. Scan en cours...\n", "success")
            results_text.update()
            ports_info = scan_ports(scanner, host, results_text)
            os_info, services_info, banners = scan_services_and_os(scanner, host, results_text)
            filename = generate_report_docx(host, ports_info, os_info, services_info, banners, folder)
            results_text.insert(tk.END, f"Rapport généré pour {host}: {filename}\n\n", "success")
        else:
            results_text.insert(tk.END, f"{host} ne répond pas. Scan annulé.\n\n", "error")
        results_text.update()
    messagebox.showinfo("Information", "Scans terminés et rapports générés.")
    results_text.insert(tk.END, "Scans terminés et rapports générés.\n", "info")

# Interface graphique Tkinter
root = tk.Tk()
root.title("Scanner de Ports et Services")
root.iconbitmap("F:/Documents/07_projet_python/programmes_ok/logos/scan_port.ico")

frame = ttk.Frame(root, padding="10")
frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

ttk.Label(frame, text="Hôtes à scanner (séparés par une virgule):").grid(row=0, column=0, sticky=tk.W, pady=5)
host_entry = ttk.Entry(frame, width=50)
host_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=5)

start_button = ttk.Button(frame, text="Démarrer le scan", command=start_scan)
start_button.grid(row=1, column=0, columnspan=2, pady=10)

results_text = scrolledtext.ScrolledText(frame, width=80, height=20, wrap=tk.WORD)
results_text.grid(row=2, column=0, columnspan=2, pady=10)

# Configuration des tags de couleur pour le texte
results_text.tag_configure("info", foreground="blue")
results_text.tag_configure("success", foreground="green")
results_text.tag_configure("warning", foreground="orange")
results_text.tag_configure("error", foreground="red")

root.mainloop()
