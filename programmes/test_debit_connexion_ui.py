import os
import subprocess
import time
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import platform
import random
import requests
from colorama import init, Fore, Style
import tkinter as tk
from tkinter import scrolledtext
from threading import Thread
from tkinter import filedialog, messagebox
import subprocess

# initialiser colorama
init()

class ConnectionSpeedTestApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Programme de Test de Connexion et de Débit")

        # ajouter l'icône
        base_dir = os.path.dirname(os.path.abspath(__file__))
        icon_path = os.path.join(base_dir, "logos", "test_connection.ico")
        self.root.iconbitmap(icon_path)

        self.root.geometry("800x600")

        self.create_widgets()
        self.all_ips = self.load_all_ips()
        self.all_domains = self.load_all_domains()
        self.ips = random.sample(self.all_ips, 20)
        self.domains = random.sample(self.all_domains, 20)

    def create_widgets(self):
        tk.Label(self.root, text="Test de Connexion et de Débit", font=("Helvetica", 16)).pack(pady=10)

        self.start_button = tk.Button(self.root, text="Démarrer les tests", command=self.start_tests)
        self.start_button.pack(pady=5)

        self.stop_button = tk.Button(self.root, text="Arrêter les tests", command=self.stop_tests, state=tk.DISABLED)
        self.stop_button.pack(pady=5)

        self.output_text = scrolledtext.ScrolledText(self.root, width=100, height=30, state=tk.DISABLED)
        self.output_text.pack(pady=10)

        self.output_text.tag_config("red", foreground="red")
        self.output_text.tag_config("green", foreground="green")
        self.output_text.tag_config("blue", foreground="blue")
        self.output_text.tag_config("magenta", foreground="magenta")

    def load_all_ips(self):
        return [
            {"ip": "8.8.8.8", "description": "DNS Public Google"},
            {"ip": "8.8.9.9", "description": "DNS Public Google (Alternatif)"},
            {"ip": "1.1.1.1", "description": "DNS Cloudflare"},
            {"ip": "1.0.0.1", "description": "DNS Cloudflare (Alternatif)"},
            {"ip": "208.67.222.222", "description": "OpenDNS"},
            {"ip": "45.90.28.0", "description": "NextDNS"},
            {"ip": "9.9.9.9", "description": "DNS Quad9"},
            {"ip": "77.88.8.8", "description": "DNS Yandex"},
            {"ip": "195.46.39.39", "description": "SafeDNS"},
            {"ip": "208.67.220.220", "description": "OpenDNS (Alternatif)"},
            {"ip": "198.41.0.4", "description": "DNS Public Verisign"},
            {"ip": "199.7.83.42", "description": "UltraDNS"},
            {"ip": "8.26.56.26", "description": "DNS Sécurisé Comodo"},
            {"ip": "8.20.247.20", "description": "DNS Sécurisé Comodo (Alternatif)"},
            {"ip": "64.6.64.6", "description": "Neustar DNS"},
            {"ip": "64.6.65.6", "description": "Neustar DNS (Alternatif)"},
            {"ip": "4.2.2.1", "description": "DNS Level3"},
            {"ip": "4.2.2.2", "description": "DNS Level3 (Alternatif)"},
            {"ip": "4.2.2.3", "description": "DNS Level3 (Alternatif)"},
            {"ip": "4.2.2.4", "description": "DNS Level3 (Alternatif)"},
            {"ip": "4.2.2.5", "description": "DNS Level3 (Alternatif)"},
            {"ip": "4.2.2.6", "description": "DNS Level3 (Alternatif)"}
        ]

    def load_all_domains(self):
        return [
            "google.fr", "google.com", "facebook.com", 
            "youtube.com", "orange.fr", "impots.gouv.fr", 
            "palamini.fr", "twitter.com", "linkedin.com", 
            "microsoft.com", "apple.com", "amazon.com", 
            "netflix.com", "instagram.com", "wikipedia.org", 
            "yahoo.com", "bing.com", "live.com", 
            "ebay.com", "blogger.com", "reddit.com", 
            "tumblr.com", "pinterest.com", "wordpress.com", 
            "whatsapp.com", "snapchat.com", "spotify.com", 
            "paypal.com", "dropbox.com", "github.com"
        ]

    def log_message(self, message, color):
        self.output_text.config(state=tk.NORMAL)
        self.output_text.insert(tk.END, message + "\n", color)
        self.output_text.config(state=tk.DISABLED)
        self.output_text.see(tk.END)

    def start_tests(self):
        report_dir = filedialog.askdirectory(title="Sélectionnez le répertoire de sauvegarde du rapport")
        if not report_dir:
            return

        self.report_dir = report_dir
        self.start_button.config(state=tk.DISABLED)
        self.stop_button.config(state=tk.NORMAL)

        thread = Thread(target=self.run_tests)
        thread.start()

    def run_tests(self):
        self.log_message("Démarrage du test de connexion...", "blue")

        ip_results = self.test_ips()
        domain_results = self.test_domains()
        download_speed = self.test_download_speed()
        upload_speed = self.test_upload_speed()

        results = {"ips": ip_results, "domains": domain_results}
        self.generate_report(results, download_speed, upload_speed, self.report_dir)
        
        # Générer un rapport parallèle dans programmes/temp
        self.generate_temp_report(results, download_speed, upload_speed)

        self.finalize_tests()

    def test_ips(self):
        ip_results = []
        self.log_message("1er test de connexion.", "magenta")
        for ip in self.ips:
            result = self.ping_ip(ip["ip"])
            ip_results.append({"ip": ip["ip"], "status": result, "description": ip["description"]})
            color = "green" if result else "red"
            self.log_message(f"{ip['ip']} ({ip['description']})... {'OK' if result else 'NON'}", color)
            time.sleep(1)
        return ip_results

    def test_domains(self):
        domain_results = []
        self.log_message("2ème test de connexion.", "magenta")
        for domain in self.domains:
            attempts = 0
            success = False
            while attempts < 3 and not success:
                attempts += 1
                success = self.ping_domain(domain) and self.nslookup(domain)
                color = "green" if success else "red"
                self.log_message(f"{domain}... {'OK' if success else f'tentative {attempts}'}", color)
                if not success:
                    time.sleep(1)
            domain_results.append({"domain": domain, "status": success, "attempts": attempts})
        return domain_results

    def test_download_speed(self):
        self.log_message("Test du débit de téléchargement...", "magenta")
        download_speed = self.perform_download_test()
        if download_speed:
            for i in range(1, 11):
                self.log_message(f"{i}...", "magenta")
                time.sleep(1)
            self.log_message("Test du débit de téléchargement terminé.", "green")
            self.log_message(f"Liaison descendante : {download_speed:.2f} Mb/s", "green")
        else:
            self.log_message("Erreur lors du test de débit de téléchargement : Impossible de se connecter aux serveurs", "red")
        return download_speed

    def test_upload_speed(self):
        self.log_message("Test du débit d'upload...", "magenta")
        upload_speed = self.perform_upload_test()
        if upload_speed:
            for i in range(1, 11):
                self.log_message(f"{i}...", "magenta")
                time.sleep(1)
            self.log_message("Test du débit d'upload terminé.", "green")
            self.log_message(f"Liaison montante : {upload_speed:.2f} Mb/s", "green")
        else:
            self.log_message("Erreur lors du test de débit d'upload : Impossible de se connecter aux serveurs", "red")
        return upload_speed

    def finalize_tests(self):
        self.start_button.config(state=tk.NORMAL)
        self.stop_button.config(state=tk.DISABLED)
        self.log_message("Tests terminés.", "green")

        # ouvrir le répertoire contenant le rapport
        if os.name == 'nt':
            os.startfile(self.report_dir)
        elif os.name == 'posix':
            subprocess.call(['xdg-open', self.report_dir])

    def stop_tests(self):
        pass

    def ping_ip(self, ip):
        param = '-n' if platform.system().lower() == 'windows' else '-c'
        command = ['ping', param, '1', ip]
        try:
            response = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=5)
            return response.returncode == 0
        except subprocess.TimeoutExpired:
            return False

    def ping_domain(self, domain):
        param = '-n' if platform.system().lower() == 'windows' else '-c'
        command = ['ping', param, '1', domain]
        try:
            response = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=5)
            return response.returncode == 0
        except subprocess.TimeoutExpired:
            return False

    def nslookup(self, domain):
        command = ['nslookup', domain]
        try:
            response = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=5)
            return response.returncode == 0
        except subprocess.TimeoutExpired:
            return False

    def perform_download_test(self):
        url = 'http://ipv4.download.thinkbroadband.com/10MB.zip'
        try:
            start_time = time.time()
            response = requests.get(url, stream=True)
            total_length = response.headers.get('content-length')

            if total_length is None:
                return None

            total_length = int(total_length)
            downloaded = 0

            for data in response.iter_content(chunk_size=4096):
                downloaded += len(data)
                if downloaded >= total_length:
                    break

            end_time = time.time()
            duration = end_time - start_time
            download_speed = (total_length * 8) / (duration * 1_000_000)
            return download_speed
        except Exception as e:
            self.log_message(f"Erreur lors du test de débit (téléchargement) : {e}", "red")
            return None

    def perform_upload_test(self):
        url = 'https://speed.hetzner.de/100MB.bin'
        data = os.urandom(10 * 1024 * 1024)
        try:
            start_time = time.time()
            response = requests.post(url, data=data)
            end_time = time.time()
            duration = end_time - start_time
            upload_speed = (len(data) * 8) / (duration * 1_000_000)
            return upload_speed
        except Exception as e:
            self.log_message(f"Erreur lors du test de débit (upload) : {e}", "red")
            return None

    def generate_report(self, results, download_speed, upload_speed, report_folder):
        date_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        date = datetime.now().strftime("%Y%m%d")
        if not os.path.exists(report_folder):
            os.makedirs(report_folder)
        document = Document()
        document.add_heading('Rapport de Connexion', 0)

        document.add_heading('Tests de Connexion IP', level=1)
        for result in results['ips']:
            color = (0, 255, 0) if result['status'] else (255, 0, 0)
            self.add_colored_paragraph(document, f"{result['ip']} ({result['description']}) : {'OK' if result['status'] else 'NON'}", color)

        document.add_heading('Tests de Connexion Domaines', level=1)
        for result in results['domains']:
            color = (0, 255, 0) if result['status'] else (255, 0, 0)
            self.add_colored_paragraph(document, f"{result['domain']} : {'OK' if result['status'] else 'NON'} (tentatives: {result['attempts']})", color)

        document.add_heading('Test de Débit', level=1)
        if download_speed is not None:
            self.add_colored_paragraph(document, f"Liaison descendante : {download_speed:.2f} Mb/s", (0, 0, 0))
        else:
            self.add_colored_paragraph(document, "Liaison descendante : Test échoué", (255, 0, 0))

        if upload_speed is not None:
            self.add_colored_paragraph(document, f"Liaison montante : {upload_speed:.2f} Mb/s", (0, 0, 0))
        else:
            self.add_colored_paragraph(document, "Liaison montante : Test échoué", (255, 0, 0))

        if download_speed is not None:
            if download_speed > 100:
                connection_type = "Fibre"
            elif download_speed > 50:
                connection_type = "VDSL"
            elif download_speed > 5:
                connection_type = "ADSL"
            else:
                connection_type = "4G/5G"
        else:
            connection_type = "Inconnue"

        document.add_heading('Qualité de la Connexion', level=1)
        self.add_colored_paragraph(document, connection_type, (0, 0, 0))

        total_ips = len(results['ips'])
        successful_ips = sum(1 for result in results['ips'] if result['status'])
        ip_success_rate = (successful_ips / total_ips) * 100

        total_domains = len(results['domains'])
        successful_domains = sum(1 for result in results['domains'] if result['status'])
        domain_success_rate = (successful_domains / total_domains) * 100

        document.add_heading('Récapitulatif Global', level=1)
        self.add_colored_paragraph(document, f"Pourcentage de réussite des tests IP : {ip_success_rate:.2f}%", (0, 0, 0))
        self.add_colored_paragraph(document, f"Pourcentage de réussite des tests de domaines : {domain_success_rate:.2f}%", (0, 0, 0))
        self.add_colored_paragraph(document, f"Nombre total d'IPs testées : {total_ips}", (0, 0, 0))
        self.add_colored_paragraph(document, f"Nombre total de domaines testés : {total_domains}", (0, 0, 0))

        report_path = os.path.join(report_folder, f"rapport_connexion_{date_time}.docx")
        document.save(report_path)

        self.log_message(f"Génération du rapport... OK\nLe rapport est sauvegardé sous : {report_path}", "green")

    def generate_temp_report(self, results, download_speed, upload_speed):
        temp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp")
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
        
        date_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        document = Document()
        document.add_heading('Rapport de Connexion (Temp)', 0)

        document.add_heading('Tests de Connexion IP', level=1)
        for result in results['ips']:
            color = (0, 255, 0) if result['status'] else (255, 0, 0)
            self.add_colored_paragraph(document, f"{result['ip']} ({result['description']}) : {'OK' if result['status'] else 'NON'}", color)

        document.add_heading('Tests de Connexion Domaines', level=1)
        for result in results['domains']:
            color = (0, 255, 0) if result['status'] else (255, 0, 0)
            self.add_colored_paragraph(document, f"{result['domain']} : {'OK' if result['status'] else 'NON'} (tentatives: {result['attempts']})", color)

        document.add_heading('Test de Débit', level=1)
        if download_speed is not None:
            self.add_colored_paragraph(document, f"Liaison descendante : {download_speed:.2f} Mb/s", (0, 0, 0))
        else:
            self.add_colored_paragraph(document, "Liaison descendante : Test échoué", (255, 0, 0))

        if upload_speed is not None:
            self.add_colored_paragraph(document, f"Liaison montante : {upload_speed:.2f} Mb/s", (0, 0, 0))
        else:
            self.add_colored_paragraph(document, "Liaison montante : Test échoué", (255, 0, 0))

        report_path = os.path.join(temp_dir, f"rapport_connexion_{date_time}.docx")
        document.save(report_path)

    def add_colored_paragraph(self, document, text, color):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])
        run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

if __name__ == "__main__":
    root = tk.Tk()
    app = ConnectionSpeedTestApp(root)
    root.mainloop()
