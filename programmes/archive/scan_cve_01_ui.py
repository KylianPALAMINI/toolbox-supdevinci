import os
import datetime
import nmap
import requests
import json
from docx import Document
from uuid import uuid4
import openai
import tkinter as tk
from tkinter import messagebox, filedialog
from tkinter import scrolledtext
from threading import Thread, Event
import subprocess

# Utilisation de votre clé API Vulners
VULNERS_API_KEY = 'RVHE0ITUYP668D4KFCXWIRZR36EFU5EX2CCLU3D4ZTBRG4YLQOC9097O65JRSP94'

# Utilisation de votre clé API de projet OpenAI
OPENAI_API_KEY = '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'

openai.api_key = OPENAI_API_KEY

class CVEScannerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Programme de Scan CVE")

        # ajouter l'icône
        base_dir = os.path.dirname(os.path.abspath(__file__))
        icon_path = os.path.join(base_dir, "logos", "scan_cve.ico")
        self.root.iconbitmap(icon_path)

        self.stop_event = Event()

        # widgets de l'interface
        self.create_widgets()

    def create_widgets(self):
        tk.Label(self.root, text="Adresse cible:").pack(pady=5)
        self.target_entry = tk.Entry(self.root, width=50)
        self.target_entry.pack(pady=5)

        self.scan_button = tk.Button(self.root, text="Démarrer le scan", command=self.start_scan)
        self.scan_button.pack(pady=5)

        self.stop_button = tk.Button(self.root, text="Arrêter le scan", command=self.stop_scan, state=tk.DISABLED)
        self.stop_button.pack(pady=5)

        self.output_text = scrolledtext.ScrolledText(self.root, width=80, height=20, state=tk.DISABLED)
        self.output_text.pack(pady=10)

        self.output_text.tag_config("red", foreground="red")
        self.output_text.tag_config("green", foreground="green")
        self.output_text.tag_config("blue", foreground="blue")

    def log_message(self, message, color):
        self.output_text.config(state=tk.NORMAL)
        self.output_text.insert(tk.END, message + "\n", color)
        self.output_text.config(state=tk.DISABLED)
        self.output_text.see(tk.END)

    def scan_services(self, target):
        self.log_message(f"Début du scan des services sur la cible : {target}", "blue")
        nm = nmap.PortScanner()
        # définir le chemin de l'exécutable Nmap
        nmap_path = r"C:\Program Files (x86)\Nmap\nmap.exe"
        os.environ["NMAP_PATH"] = nmap_path
        nm.scan(target, arguments='-sV -n')
        services = []
        for host in nm.all_hosts():
            for proto in nm[host].all_protocols():
                lport = nm[host][proto].keys()
                for port in lport:
                    service = {
                        'name': nm[host][proto][port]['name'],
                        'version': nm[host][proto][port].get('version', 'inconnue')
                    }
                    services.append(service)
        self.log_message("Scan des services terminé", "blue")
        return services

    def get_cve(self, service_name, service_version):
        url = 'https://vulners.com/api/v3/search/lucene/'
        headers = {
            'Content-Type': 'application/json',
            'X-Vulners-Api-Key': VULNERS_API_KEY
        }
        query = f"{service_name} {service_version}"
        data = {
            "query": query,
            "size": 50
        }

        try:
            response = requests.post(url, headers=headers, data=json.dumps(data))
            response.raise_for_status()
            try:
                cve_data = response.json()
                if 'data' in cve_data and 'search' in cve_data['data']:
                    documents = cve_data['data']['search']
                    return [
                        {
                            "id": doc['_id'],
                            "summary": doc.get('description', doc.get('title', 'Pas de description disponible'))
                        } for doc in documents if doc['_id'].startswith('CVE-')
                    ]
                return []
            except json.JSONDecodeError as e:
                self.log_message(f"Erreur de décodage JSON : {e}", "red")
                self.log_message("Réponse reçue de l'API : " + response.text, "red")
                return []
        except requests.RequestException as e:
            self.log_message(f"Erreur lors de la requête HTTP : {e}", "red")
            return []

    def generate_service_description(self, service_name):
        prompt = f"Fournissez une description concise en français pour le service {service_name} en expliquant à quoi il sert."
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Vous êtes un expert en cybersécurité."},
                    {"role": "user", "content": prompt}
                ]
            )
            description = response.choices[0].message['content'].strip()

            if "Désolé" in description in "ne peux pas fournir" in description:
                return "Pas de description disponible"

            return description
        except openai.error.OpenAIError as e:
            self.log_message(f"Erreur lors de la génération de la description du service {service_name} : {e}", "red")
            return "Pas de description disponible"

    def generate_cve_description(self, cve_id):
        prompt = f"Fournissez une description concise pour le CVE {cve_id} en français, en mentionnant le type de vulnérabilité, les systèmes affectés et l'impact potentiel. Limitez la description à 3 phrases."
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Vous êtes un expert en cybersécurité."},
                    {"role": "user", "content": prompt}
                ]
            )
            description = response.choices[0].message['content'].strip()

            if "Désolé" in description in "ne peux pas fournir" in description:
                return "Pas de description disponible"

            return description
        except openai.error.OpenAIError as e:
            self.log_message(f"Erreur lors de la génération de la description pour {cve_id} : {e}", "red")
            return "Pas de description disponible"

    def generate_cve_recommendation(self, cve_id):
        prompt = f"Fournissez une recommandation concise pour le CVE {cve_id} en français, incluant les étapes pour mitiger la vulnérabilité. Limitez la recommandation à 3 phrases."
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Vous êtes un expert en cybersécurité."},
                    {"role": "user", "content": prompt}
                ]
            )
            recommendation = response.choices[0].message['content'].strip()

            if "Désolé" in recommendation in "ne peux pas fournir" in recommendation:
                return "Pas de recommandation disponible"

            return recommendation
        except openai.error.OpenAIError as e:
            self.log_message(f"Erreur lors de la génération de la recommandation pour {cve_id} : {e}", "red")
            return "Pas de recommandation disponible"

    def generate_report(self, services, cve_data, base_dir):
        if not os.path.exists(base_dir):
            os.makedirs(base_dir)

        today = datetime.datetime.now().strftime("%Y%m%d")
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = uuid4()
        report_dir = os.path.join(base_dir, f"rapport_scan_cve_{today}")
        if not os.path.exists(report_dir):
            os.makedirs(report_dir)
        report_name = f"scan_cve_{timestamp}_{unique_id}.docx"
        report_path = os.path.join(report_dir, report_name)

        doc = Document()
        doc.add_heading('Rapport de Scan CVE', 0)

        doc.add_heading('Services', level=1)
        for service in services:
            service_description = self.generate_service_description(service['name'])
            doc.add_paragraph(f"Service: {service['name']}, Version: {service['version']}")
            doc.add_paragraph(f"Description: {service_description}")

        doc.add_heading('Vulnérabilités', level=1)
        for service in services:
            cve_list = cve_data.get(service['name'], [])
            if cve_list:
                doc.add_heading(f"Service: {service['name']} (Version: {service['version']})", level=2)
                for cve in cve_list:
                    cve_description = self.generate_cve_description(cve['id'])
                    cve_recommendation = self.generate_cve_recommendation(cve['id'])
                    p = doc.add_paragraph()
                    p.add_run(f"CVE ID: {cve['id']}\n").bold = True
                    p.add_run(f"Description: {cve_description}\n")
                    p.add_run(f"Recommandation: {cve_recommendation}\n").italic = True

        doc.save(report_path)
        self.log_message(f"Rapport généré : {report_path}", "green")
        return report_path

    def start_scan(self):
        self.clear_output()
        target = self.target_entry.get()
        if not target:
            messagebox.showerror("Erreur", "Veuillez entrer l'adresse cible.")
            return

        base_dir = os.path.dirname(os.path.abspath(__file__))

        self.scan_button.config(state=tk.DISABLED)
        self.stop_button.config(state=tk.NORMAL)

        self.stop_event.clear()
        thread = Thread(target=self.scan_and_generate_report, args=(target, base_dir))
        thread.start()

    def scan_and_generate_report(self, target, base_dir):
        self.log_message(f"Démarrage du scan pour la cible : {target}", "blue")
        services = self.scan_services(target)
        if self.stop_event.is_set():
            self.log_message("Le scan a été arrêté.", "red")
            return
        self.log_message("Services scannés : " + str(services), "blue")

        cve_data = {}
        for service in services:
            if self.stop_event.is_set():
                self.log_message("Le scan a été arrêté.", "red")
                return
            cve_data[service['name']] = self.get_cve(service['name'], service['version'])
            self.log_message(f"CVE pour {service['name']} : {cve_data[service['name']]}", "blue")

        if not self.stop_event.is_set():
            self.log_message("Génération du rapport...", "blue")
            report_path = self.generate_report(services, cve_data, base_dir)
            self.finalize_scan(report_path)
        else:
            self.log_message("Le scan a été arrêté avant la génération du rapport.", "red")

    def finalize_scan(self, report_path):
        self.scan_button.config(state=tk.NORMAL)
        self.stop_button.config(state=tk.DISABLED)
        self.log_message("Scan terminé.", "green")

        # ouvrir le répertoire contenant le rapport
        report_dir = os.path.dirname(report_path)
        if os.name == 'nt':
            os.startfile(report_dir)
        elif os.name == 'posix':
            subprocess.call(['xdg-open', report_dir])

    def stop_scan(self):
        self.stop_event.set()
        self.log_message("Arrêt du scan en cours...", "red")
        self.scan_button.config(state=tk.NORMAL)
        self.stop_button.config(state=tk.DISABLED)

    def clear_output(self):
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete('1.0', tk.END)
        self.output_text.config(state=tk.DISABLED)

if __name__ == "__main__":
    root = tk.Tk()
    app = CVEScannerApp(root)
    root.mainloop()
