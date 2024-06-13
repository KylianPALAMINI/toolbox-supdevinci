import os
import shodan
import sys
from scapy.all import *
from datetime import datetime
from colorama import init, Fore, Style
from docx import Document
from docx.shared import Pt
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext

# Initialiser Colorama
init(autoreset=True)

# Configuration de l'API Shodan
SHODAN_API_KEY = 'H62LeLAdwcuguJsSfeIgMPB2lNatohFW'
api = shodan.Shodan(SHODAN_API_KEY)

def verifier_cle_api(text_widget):
    try:
        # Vérifier l'accès à l'API Shodan
        info = api.info()
        text_widget.insert(tk.END, "Accès à l'API Shodan vérifié avec succès.\n", "success")
        text_widget.insert(tk.END, "Informations sur l'API Shodan:\n", "info")
        text_widget.insert(tk.END, f"{info}\n", "info")
        if info['query_credits'] > 0:
            return True
        else:
            text_widget.insert(tk.END, "Vous n'avez pas assez de crédits de requête Shodan.\n", "error")
            return False
    except shodan.APIError as e:
        text_widget.insert(tk.END, f"Erreur lors de l'accès à l'API Shodan: {e}\n", "error")
        return False

def recherche_shodan(cible, text_widget):
    try:
        # Recherche d'informations sur la cible avec Shodan
        host = api.host(cible)
        results = []
        results.append(f"IP: {host['ip_str']}")
        results.append(f"Organisation: {host.get('org', 'n/a')}")
        results.append(f"Système d'exploitation: {host.get('os', 'n/a')}")
        
        # Affichage des ports ouverts et des services
        for item in host['data']:
            product = item.get('product', 'n/a')
            version = item.get('version', 'n/a')
            port_info = f"Port: {item['port']} - Service: {product} {version}"
            results.append(port_info)
            results.append(f"Bannière: {item['data']}")
            
            # Affichage des vulnérabilités trouvées
            if 'vulns' in item:
                for vuln in item['vulns']:
                    cve = vuln.replace('!', '')
                    vuln_info = f"Vulnérabilité: {vuln}"
                    results.append(vuln_info)
                    if 'cve' in item['vulns'][vuln]:
                        results.append(f"CVE: {item['vulns'][vuln]['cve']}")
                    if 'cvss' in item['vulns'][vuln]:
                        results.append(f"Score CVSS: {item['vulns'][vuln]['cvss']}")
                    if 'summary' in item['vulns'][vuln]:
                        results.append(f"Résumé: {item['vulns'][vuln]['summary']}")
        return results
    except shodan.APIError as e:
        text_widget.insert(tk.END, f"Erreur: {e}\n", "error")
        return []

def scan_scapy(cible, text_widget):
    # Scan des ports avec Scapy
    text_widget.insert(tk.END, f"Démarrage du scan Scapy sur {cible}...\n", "info")
    open_ports = []
    try:
        ans, unans = sr(IP(dst=cible)/TCP(dport=(1,1024),flags="S"),timeout=5, verbose=False)
        for sent, received in ans:
            if received.haslayer(TCP) and received.getlayer(TCP).flags == 0x12:
                port_info = f"Le port {sent.dport} est ouvert"
                open_ports.append(port_info)
                text_widget.insert(tk.END, f"{port_info}\n", "success")
    except Exception as e:
        text_widget.insert(tk.END, f"Une erreur est survenue lors du scan Scapy: {e}\n", "error")
    return open_ports

def creer_rapport(date_str, cible, shodan_results, scapy_results):
    # Créer le dossier daté
    dossier = f"rapport_scan_shodan_{date_str}"
    if not os.path.exists(dossier):
        os.makedirs(dossier)
    
    # Obtenir l'heure actuelle pour le nom de fichier unique
    heure_str = datetime.now().strftime("%H-%M-%S")
    
    # Créer le document Word
    doc = Document()
    doc.add_heading(f'Rapport de Scan de Vulnérabilités - {date_str}', 0)

    # Ajouter les résultats Shodan
    doc.add_heading('Résultats Shodan', level=1)
    for line in shodan_results:
        p = doc.add_paragraph(line)
        p.style.font.size = Pt(12)
    
    # Ajouter les résultats Scapy
    doc.add_heading('Résultats Scapy', level=1)
    for line in scapy_results:
        p = doc.add_paragraph(line)
        p.style.font.size = Pt(12)
    
    # Sauvegarder le rapport
    rapport_path = os.path.join(dossier, f'rapport_{cible}_{date_str}_{heure_str}.docx')
    doc.save(rapport_path)
    return rapport_path

def start_scan():
    cible = cible_entry.get()
    if not cible:
        messagebox.showwarning("Avertissement", "Veuillez entrer l'adresse IP ou le nom d'hôte de la cible.")
        return
    
    results_text.delete(1.0, tk.END)
    date_str = datetime.now().strftime("%d-%m-%Y")
    
    if verifier_cle_api(results_text):
        shodan_results = recherche_shodan(cible, results_text)
    else:
        shodan_results = []

    scapy_results = scan_scapy(cible, results_text)
    rapport_path = creer_rapport(date_str, cible, shodan_results, scapy_results)
    results_text.insert(tk.END, f"Rapport généré et sauvegardé sous {rapport_path}\n", "info")
    messagebox.showinfo("Information", "Scans terminés et rapport généré.")
    results_text.insert(tk.END, "Scans terminés et rapport généré.\n", "success")

# Interface graphique Tkinter
class VulnScannerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Scanner de Vulnérabilités")

        # Ajouter le logo
        icon_path = "F:/Documents/07_projet_python/programmes_ok/logos/scan_vuln.ico"
        self.root.iconbitmap(icon_path)

        frame = ttk.Frame(root, padding="10")
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        ttk.Label(frame, text="Adresse IP ou nom d'hôte de la cible:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.cible_entry = ttk.Entry(frame, width=50)
        self.cible_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=5)

        self.start_button = ttk.Button(frame, text="Démarrer le scan", command=start_scan)
        self.start_button.grid(row=1, column=0, columnspan=2, pady=10)

        self.results_text = scrolledtext.ScrolledText(frame, width=80, height=20, wrap=tk.WORD)
        self.results_text.grid(row=2, column=0, columnspan=2, pady=10)

        # Configuration des tags de couleur pour le texte
        self.results_text.tag_configure("info", foreground="blue")
        self.results_text.tag_configure("success", foreground="green")
        self.results_text.tag_configure("warning", foreground="orange")
        self.results_text.tag_configure("error", foreground="red")

if __name__ == "__main__":
    root = tk.Tk()
    app = VulnScannerApp(root)
    root.mainloop()
